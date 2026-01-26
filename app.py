import streamlit as st
from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
from datetime import datetime

# Carregar chave API
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

st.set_page_config(page_title="Agente Monografia", layout="wide")
st.title("🎓 Sistema Agente Monografia")

# --- Inicialização da Memória dos Agentes ---
if "step" not in st.session_state:
    st.session_state.step = 1
if "dados" not in st.session_state:
    st.session_state.dados = {}

def call_gpt(prompt):
    response = client.chat.completions.create(
        model="gpt-4o", # ou gpt-3.5-turbo
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# --- AGENTE 1: ESCOLHA DO TEMA ---
if st.session_state.step == 1:
    st.header("Agente 1: Escolha do Tema")
    area = st.text_input("Qual sua área de interesse?")
    
    if st.button("Analisar Tendências"):
        prompt = f"""Usando o google search, faça uma análise dos títulos dos artigos dos últimos 5
          anos dos 3 periódicos científicos internacionais mais importantes atualmente na área {area},
          identifique os 5 temas mais relevantes com base na quantidade de publicações e apresente-os como sugestões
          de temas para a monografia do usuário. Escreva na tela apenas as sugestões. 
        """
        resultado = call_gpt(prompt)
        st.session_state.temas_sugeridos = resultado
        st.rerun()

    if "temas_sugeridos" in st.session_state:
        st.markdown(st.session_state.temas_sugeridos)
        escolha = st.text_input("Digite o tema escolhido ou, caso não tenha gostado de nenhuma das minhas sugestões, um novo tema:")
        if st.button("Confirmar Tema"):
            st.session_state.dados['tema_base'] = escolha
            st.session_state.step = 2
            st.rerun()

# --- AGENTE 2: APROFUNDAMENTO ---
elif st.session_state.step == 2:
    st.header("Agente 2: Aprofundamento do Tema")
    if "subtemas" not in st.session_state:
        prompt = f"""O usuário demonstrou interesse no tema {st.session_state.dados['tema_base']}. 
        Apresente ao usuário 10 sugestões de subtemas mais específicos para a monografia dele. 
        Considere sempre que a monografia será uma revisão da literatura."""
        st.session_state.subtemas = call_gpt(prompt)

    st.write(f"**Tema macro:** {st.session_state.dados['tema_base']}")
    st.markdown(st.session_state.subtemas)
    
    # Extrair linhas para o selectbox (simulação simplificada)
    linhas = st.session_state.subtemas.split('\n')
    opcoes = [st.session_state.dados['tema_base']] + [l for l in linhas if l.strip() and l[0].isdigit()]
    
    escolha_sub = st.radio("Marque qual tema devemos trabalhar:", opcoes)
    
    if st.button("Confirmar Subtema"):
        st.session_state.dados['tema_escolhido'] = escolha_sub
        st.session_state.step = 3
        st.rerun()

# --- AGENTE 3: PROBLEMA DE PESQUISA ---
elif st.session_state.step == 3:
    st.header("Agente 3: Problema de Pesquisa")
    if "problemas" not in st.session_state:
        prompt = f"Quero escrever uma monografia que será uma revisão da literatura sobre o tema {st.session_state.dados['tema_escolhido']}. Crie 5 sugestões de 'Problema de pesquisa'."
        st.session_state.problemas = call_gpt(prompt)

    st.write(f"**Tema Escolhido:** {st.session_state.dados['tema_escolhido']}")
    st.markdown(st.session_state.problemas)
    
    problema_final = st.text_area("Escolha um dos problemas digitando o número ou digite seu próprio:")
    
    if st.button("Confirmar Problema"):
        st.session_state.dados['problema_pesquisa'] = problema_final
        st.session_state.step = 4
        st.rerun()

# --- AGENTE 4: OBJETIVOS ESPECÍFICOS ---
elif st.session_state.step == 4:
    st.header("Agente 4: Objetivos Específicos")
    if "objetivos_lista" not in st.session_state:
        prompt = f"Considerando que o tema será {st.session_state.dados['tema_escolhido']} e o Problema de Pesquisa {st.session_state.dados['problema_pesquisa']}, crie 10 sugestões de objetivos específicos."
        res = call_gpt(prompt)
        st.session_state.objetivos_lista = [l for l in res.split('\n') if l.strip() and l[0].isdigit()]

    selecionados = []
    for obj in st.session_state.objetivos_lista:
        if st.checkbox(obj, key=obj):
            selecionados.append(obj)
            
    if st.button("Confirmar Objetivos"):
        st.session_state.dados['objetivos'] = "\n".join(selecionados)
        # Gatilho para Agentes 5 e 6
        st.session_state.step = 5
        st.rerun()

# --- AGENTE 5 & 6: REFERÊNCIAS ---
elif st.session_state.step == 5:
    st.header("Processando Referências (Agentes 5 e 6)...")
    
    with st.spinner("Buscando Referências Clássicas e Atuais..."):
        # Agente 5
        p5 = f"Revisão de literatura sobre {st.session_state.dados['tema_escolhido']}, problema {st.session_state.dados['problema_pesquisa']}, objetivos {st.session_state.dados['objetivos']}. Quais trabalhos clássicos e autores importantes?"
        ref_classicas = call_gpt(p5)
        
        # Agente 6
        ano_atual = datetime.now().year
        p6 = f"Atue como especialista em metodologia... Tema: {st.session_state.dados['tema_escolhido']}, Problema: {st.session_state.dados['problema_pesquisa']}, Objetivos: {st.session_state.dados['objetivos']}. Busca entre {ano_atual} e {ano_atual-5}..."
        ref_atuais = call_gpt(p6)
        
        st.session_state.dados['ref_classicas'] = ref_classicas
        st.session_state.dados['ref_atuais'] = ref_atuais
        st.session_state.step = 6
        st.rerun()

# --- AGENTE 7: CONSOLIDAÇÃO E DOWNLOAD ---
elif st.session_state.step == 6:
    st.header("Agente 7: Consolidação")
    
    doc = Document()
    doc.add_heading('Estrutura de Monografia Gerada por IA', 0)
    
    for chave, valor in st.session_state.dados.items():
        doc.add_heading(chave.replace('_', ' ').title(), level=1)
        doc.add_paragraph(str(valor))
    
    bio = BytesIO()
    doc.save(bio)
    
    st.success("Documento consolidado com sucesso!")
    st.download_button(
        label="Download Monografia (.docx)",
        data=bio.getvalue(),
        file_name="monografia_projeto.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    if st.button("Reiniciar Processo"):
        st.session_state.clear()
        st.rerun()